import types
import asyncio
from io import BytesIO
import openpyxl
import docx
import mammoth
import os
import pytest

import bot.settings as settings
from bot.bot import OpenAIBot, setup_openai, TelegramBot


@pytest.fixture(autouse=True)
def _patch_openai(monkeypatch):
    models_resp = types.SimpleNamespace(
        data=[types.SimpleNamespace(id="m1"), types.SimpleNamespace(id="m2")]
    )
    fake_openai = types.SimpleNamespace(
        api_key=None,
        Model=types.SimpleNamespace(list=lambda: models_resp),
        chat=types.SimpleNamespace(
            completions=types.SimpleNamespace(create=lambda **k: None)
        ),
    )
    monkeypatch.setattr("bot.bot.openai", fake_openai)


def test_setup_openai(monkeypatch):
    fake_openai = types.SimpleNamespace(api_key=None)
    monkeypatch.setattr(settings, "OPENAI_API_KEY", "key")
    monkeypatch.setattr("bot.bot.openai", fake_openai)
    setup_openai()
    assert fake_openai.api_key == "key"


def test_openai_bot_sequential(monkeypatch):
    assistants = [
        {"role": "assistant1", "system_prompt": "A"},
        {"role": "assistant2", "system_prompt": "B"},
    ]
    monkeypatch.setattr(settings, "ASSISTANTS", assistants)

    responses = [
        {"choices": [{"message": {"content": "first"}}]},
        {"choices": [{"message": {"content": "second"}}]},
    ]

    def fake_create(**kwargs):
        return responses.pop(0)

    monkeypatch.setattr("bot.bot._create_chat_completion", fake_create)
    bot_instance = OpenAIBot()
    reply = list(bot_instance.ask_stream("hello"))
    assert reply == [("assistant1", "first"), ("assistant2", "second")]


def test_openai_bot_image(monkeypatch):
    content = [
        {"type": "text", "text": "hi"},
        {"type": "image_url", "image_url": {"url": "data:image/png;base64,abc"}},
    ]

    def fake_create(**kwargs):
        assert kwargs["messages"][1]["content"] == content
        return {"choices": [{"message": {"content": "ok"}}]}

    monkeypatch.setattr("bot.bot._create_chat_completion", fake_create)
    bot_instance = OpenAIBot()
    reply = bot_instance.ask(content)
    assert reply == "ok"



def test_openai_bot_conversation(monkeypatch):
    monkeypatch.setattr(
        settings, "ASSISTANTS", [{"role": "assistant", "system_prompt": "A"}]
    )

    replies = [
        {"choices": [{"message": {"content": "first"}}]},
        {"choices": [{"message": {"content": "second"}}]},
    ]

    def fake_create(**kwargs):
        return replies.pop(0)

    monkeypatch.setattr("bot.bot._create_chat_completion", fake_create)
    bot_instance = OpenAIBot()
    conv = []
    bot_instance.ask("hi", conv)
    bot_instance.ask("again", conv)

    assert conv == [
        {"role": "user", "content": "hi"},
        {"role": "assistant", "content": "first"},
        {"role": "user", "content": "again"},
        {"role": "assistant", "content": "second"},
    ]

def test_telegram_bot_album(monkeypatch):
    async def run():
        responses = []

        class DummyPhoto:
            async def get_file(self):
                class F:
                    async def download_as_bytearray(self):
                        return b"img"

                return F()

        class DummyMessage:
            def __init__(self, caption=None):
                self.caption = caption
                self.text = None
                self.photo = [DummyPhoto()]
                self.document = None
                self.audio = None
                self.media_group_id = "g"
                self.chat_id = 1

        bot_instance = TelegramBot()
        monkeypatch.setattr(
            bot_instance.bot, "ask_stream", lambda c: iter([("assistant", "album")])
        )

        async def send_message(chat_id, text):
            responses.append(text)

        ctx = types.SimpleNamespace(bot=types.SimpleNamespace(send_message=send_message))
        bot_instance._media_groups["g"] = {"messages": [DummyMessage("hi"), DummyMessage()], "task": None}
        await bot_instance._process_media_group("g", ctx)

        assert responses == ["album"]

    asyncio.run(run())


def test_telegram_bot_document_group(monkeypatch):
    async def run():
        responses = []

        class DummyDocument:
            file_name = "file.pdf"

        class DummyMessage:
            def __init__(self, caption=None):
                self.caption = caption
                self.text = None
                self.photo = None
                self.document = DummyDocument()
                self.audio = None
                self.media_group_id = "g"
                self.chat_id = 1

        bot_instance = TelegramBot()
        async def dummy_read(d):
            return "doc", []
        monkeypatch.setattr(bot_instance, "_read_document_text", dummy_read)
        monkeypatch.setattr(
            bot_instance.bot, "ask_stream", lambda c: iter([("assistant", "docs")])
        )

        async def send_message(chat_id, text):
            responses.append(text)

        ctx = types.SimpleNamespace(bot=types.SimpleNamespace(send_message=send_message))
        bot_instance._media_groups["g"] = {"messages": [DummyMessage("hi"), DummyMessage()], "task": None}
        await bot_instance._process_media_group("g", ctx)

        assert responses == ["docs"]

    asyncio.run(run())


def test_telegram_bot_single_document(monkeypatch):
    async def run():
        run.reply = None
        class DummyDocument:
            file_name = "report.pdf"

        class DummyMessage:
            caption = None
            text = None
            photo = None
            document = DummyDocument()
            audio = None
            media_group_id = None

            def __init__(self):
                self.chat_id = 1

            async def reply_text(self, text):
                run.reply = text

        bot_instance = TelegramBot()
        async def dummy_read(d):
            return "some text", []
        monkeypatch.setattr(bot_instance, "_read_document_text", dummy_read)

        def fake_stream(content, conv):
            assert "some text" in content
            return iter([("assistant", "ok")])

        monkeypatch.setattr(bot_instance.bot, "ask_stream", fake_stream)
        update = types.SimpleNamespace(message=DummyMessage(), effective_chat=types.SimpleNamespace(id=1))
        ctx = types.SimpleNamespace(application=types.SimpleNamespace(create_task=lambda c: None))
        await bot_instance.handle_message(update, ctx)
        assert run.reply == "ok"

    asyncio.run(run())


def _dummy_document(name: str, data: bytes):
    class Dummy:
        file_name = name

        async def get_file(self):
            data_bytes = data

            class F:
                async def download_as_bytearray(self):
                    return data_bytes

            return F()

    return Dummy()


def test_read_document_text_csv():
    async def run():
        bot_instance = TelegramBot()
        doc = _dummy_document("data.csv", b"a,b\n1,2")
        text, images = await bot_instance._read_document_text(doc)
        assert "1,2" in text
        assert images == []

    asyncio.run(run())


def test_read_document_text_xlsx():
    async def run():
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "hello"
        bio = BytesIO()
        wb.save(bio)
        doc = _dummy_document("file.xlsx", bio.getvalue())
        bot_instance = TelegramBot()
        text, images = await bot_instance._read_document_text(doc)
        assert "hello" in text
        assert images == []

    asyncio.run(run())


def test_read_document_text_docx():
    async def run():
        docx_file = docx.Document()
        docx_file.add_paragraph("test")
        table = docx_file.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "cell"
        # Add dummy image
        from PIL import Image
        from io import BytesIO
        img = Image.new("RGB", (10, 10), color="red")
        img_bio = BytesIO()
        img.save(img_bio, format="PNG")
        img_path = "test.png"
        with open(img_path, "wb") as f:
            f.write(img_bio.getvalue())
        docx_file.add_picture(img_path)
        bio = BytesIO()
        docx_file.save(bio)
        bot_instance = TelegramBot()
        doc = _dummy_document("doc.docx", bio.getvalue())
        text, images = await bot_instance._read_document_text(doc)
        assert "test" in text
        assert "cell" in text
        assert images and images[0].startswith("data:image")
        os.remove(img_path)

    asyncio.run(run())


def test_read_document_text_doc():
    async def run():
        bot_instance = TelegramBot()
        data = "doc text".encode("utf-16le")
        doc = _dummy_document("sample.doc", data)
        text, images = await bot_instance._read_document_text(doc)
        assert "doc text" in text
        assert images == []

    asyncio.run(run())


def test_consilium_mode(monkeypatch):
    async def run():
        responses = []

        class DummyMessage:
            caption = None
            text = "hi"
            photo = None
            document = None
            audio = None
            media_group_id = None

            def __init__(self):
                self.chat_id = 1

            async def reply_text(self, text):
                responses.append(text)

        bot_instance = TelegramBot()

        # default bot reply
        monkeypatch.setattr(
            bot_instance.bot, "ask_stream", lambda c, conv=None: iter([("assistant", "default")])
        )

        update = types.SimpleNamespace(message=DummyMessage(), effective_chat=types.SimpleNamespace(id=1))
        ctx = types.SimpleNamespace(application=types.SimpleNamespace(create_task=lambda c: None))

        await bot_instance.start_consilium(update, ctx)
        assert 1 in bot_instance.bots

        monkeypatch.setattr(
            bot_instance.bots[1], "ask_stream", lambda c, conv=None: iter([("doctor", "consilium")])
        )
        await bot_instance.handle_message(update, ctx)
        assert responses[-1] == "doctor: consilium"

        await bot_instance.stop_consilium(update, ctx)
        assert 1 not in bot_instance.bots
        await bot_instance.handle_message(update, ctx)
        assert responses[-1] == "default"

    asyncio.run(run())


def test_model_selection(monkeypatch):
    async def run():
        responses = []

        bot_instance = TelegramBot()

        class DummyMessage:
            async def reply_text(self, text):
                responses.append(text)

        update = types.SimpleNamespace(
            message=DummyMessage(), effective_chat=types.SimpleNamespace(id=1)
        )
        ctx = types.SimpleNamespace(args=["m2"])
        await bot_instance.set_model(update, ctx)

        assert bot_instance.bot.model == "m2"
        assert responses[-1].startswith("Model set to m2")

    asyncio.run(run())


def test_bot_selection(monkeypatch):
    async def run():
        responses = []

        bot_instance = TelegramBot()

        class DummyMessage:
            async def reply_text(self, text):
                responses.append(text)

        update = types.SimpleNamespace(
            message=DummyMessage(), effective_chat=types.SimpleNamespace(id=1)
        )
        ctx = types.SimpleNamespace(args=["claude"])
        await bot_instance.set_bot(update, ctx)

        assert bot_instance.bot_name == "claude"
        assert responses[-1].startswith("Bot set to claude")

    asyncio.run(run())

